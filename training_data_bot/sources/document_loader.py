"""Document loader for text files (txt, md, docx, etc.)."""

import os
from pathlib import Path
from typing import Any

import aiofiles
import docx

from training_data_bot.sources.base import BaseLoader


class DocumentLoader(BaseLoader):
    """Loader for text-based document files."""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".json", ".xml", ".csv", ".log"}

    async def load_single(self, source: str) -> dict[str, Any]:
        """
        Load a single document file.

        Args:
            source: Path to the document file

        Returns:
            Dictionary containing source, content, and metadata
        """
        file_path = Path(source)

        if not file_path.exists():
            self.logger.warning(f"File not found: {source}")
            return {
                "source": source,
                "doc_type": "document",
                "content": "",
                "status": "error",
                "error": f"File not found: {source}",
            }

        extension = file_path.suffix.lower()

        if extension == ".docx":
            return await self._load_docx(file_path, source)
        elif extension in self.SUPPORTED_EXTENSIONS:
            return await self._load_text_file(file_path, source)
        else:
            self.logger.warning(f"Unsupported file extension: {extension}")
            return {
                "source": source,
                "doc_type": "document",
                "content": "",
                "status": "error",
                "error": f"Unsupported file extension: {extension}",
            }

    async def _load_text_file(self, file_path: Path, source: str) -> dict[str, Any]:
        """Load a plain text file asynchronously."""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()

            file_stat = os.stat(file_path)

            return {
                "source": source,
                "doc_type": "document",
                "content": content,
                "status": "success",
                "metadata": {
                    "file_name": file_path.name,
                    "file_size": file_stat.st_size,
                    "extension": file_path.suffix.lower(),
                    "encoding": "utf-8",
                    "line_count": len(content.splitlines()),
                    "char_count": len(content),
                },
            }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, "r", encoding="latin-1") as f:
                    content = await f.read()

                file_stat = os.stat(file_path)

                return {
                    "source": source,
                    "doc_type": "document",
                    "content": content,
                    "status": "success",
                    "metadata": {
                        "file_name": file_path.name,
                        "file_size": file_stat.st_size,
                        "extension": file_path.suffix.lower(),
                        "encoding": "latin-1",
                        "line_count": len(content.splitlines()),
                        "char_count": len(content),
                    },
                }
            except Exception as e:
                self.logger.error(f"Error reading file {source}: {e}")
                return {
                    "source": source,
                    "doc_type": "document",
                    "content": "",
                    "status": "error",
                    "error": str(e),
                }
        except Exception as e:
            self.logger.error(f"Error reading file {source}: {e}")
            return {
                "source": source,
                "doc_type": "document",
                "content": "",
                "status": "error",
                "error": str(e),
            }

    async def _load_docx(self, file_path: Path, source: str) -> dict[str, Any]:
        """Load a DOCX file."""
        try:
            doc = docx.Document(file_path)

            # Extract text from all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n".join(paragraphs)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content += "\n" + cell.text

            file_stat = os.stat(file_path)

            return {
                "source": source,
                "doc_type": "document",
                "content": content,
                "status": "success",
                "metadata": {
                    "file_name": file_path.name,
                    "file_size": file_stat.st_size,
                    "extension": ".docx",
                    "encoding": "binary",
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "char_count": len(content),
                },
            }
        except Exception as e:
            self.logger.error(f"Error reading DOCX file {source}: {e}")
            return {
                "source": source,
                "doc_type": "document",
                "content": "",
                "status": "error",
                "error": str(e),
            }
